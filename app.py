import streamlit as st
from pathlib import Path
from datetime import datetime
import re
from typing import Optional
from xml.sax.saxutils import escape
import os

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import mm
from reportlab.lib.enums import TA_LEFT

from docx import Document
from pypdf import PdfReader

try:
    from backend.transcriber import extract_audio, transcribe_audio
except ImportError:
    extract_audio = None
    transcribe_audio = None

try:
    from backend.summarizer import summarize_text
except ImportError:
    summarize_text = None

try:
    from backend.utils import chunked_summarize
except ImportError:
    chunked_summarize = None

try:
    from backend.pipeline import run_feedback_pipeline
except ImportError:
    run_feedback_pipeline = None

try:
    from backend.presentation_evaluator import evaluate_presentation
except ImportError:
    evaluate_presentation = None

try:
    from backend.mentor_evaluator import generate_mentor_feedback_pack
except ImportError:
    generate_mentor_feedback_pack = None


BRIEF_FOLDER = Path("data/briefs")
SUBMISSION_FOLDER = Path("data/submissions")
PRESENTATION_FOLDER = Path("data/presentation_transcripts")
REPORT_FOLDER = Path("data/reports")


def apply_custom_css():
    st.markdown(
        """
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@400;500;600;700&display=swap');

        html, body, [class*="css"], .stApp {
            font-family: 'Montserrat', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
            font-size: 14px !important;
        }

        .stApp {
            background: #ffffff;
            color: #1e293b;
        }

        [data-testid="stHeader"] {
            background: rgba(255, 255, 255, 0.85) !important;
            backdrop-filter: blur(18px) saturate(1.2);
            -webkit-backdrop-filter: blur(18px) saturate(1.2);
            border-bottom: 1px solid rgba(26, 38, 58, 0.1);
        }

        h1 {
            color: #1a3a52 !important;
            font-family: 'Montserrat', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
            font-weight: 600 !important;
            font-size: 1.6rem !important;
            line-height: 1.3 !important;
            letter-spacing: -0.3px !important;
        }

        h2 {
            color: #1a3a52 !important;
            font-family: 'Montserrat', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
            font-weight: 600 !important;
            font-size: 1.25rem !important;
            line-height: 1.4 !important;
        }

        h3 {
            color: #c9a227 !important;
            font-family: 'Montserrat', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
            font-weight: 600 !important;
            font-size: 1.05rem !important;
            line-height: 1.4 !important;
        }

        h4, h5, h6 {
            color: #c9a227 !important;
            font-family: 'Montserrat', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
            font-weight: 500 !important;
            font-size: 0.95rem !important;
        }

        [data-testid="stMarkdownContainer"] h1 {
            color: #1a3a52 !important;
            font-family: 'Montserrat', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
            font-weight: 600 !important;
            font-size: 1.6rem !important;
        }

        [data-testid="stMarkdownContainer"] h2 {
            color: #1a3a52 !important;
            font-family: 'Montserrat', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
            font-weight: 600 !important;
            font-size: 1.25rem !important;
        }

        [data-testid="stMarkdownContainer"] h3 {
            color: #c9a227 !important;
            font-family: 'Montserrat', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
            font-weight: 600 !important;
            font-size: 1.05rem !important;
        }

        p, label, span, div {
            color: #374151 !important;
            font-family: 'Montserrat', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
            font-size: 14px !important;
            line-height: 1.6 !important;
        }

        /* --- Glassmorphism Cards --- */
        .custom-card {
            background: rgba(255, 255, 255, 0.85);
            backdrop-filter: blur(16px) saturate(1.3);
            -webkit-backdrop-filter: blur(16px) saturate(1.3);
            border: 1px solid rgba(26, 38, 58, 0.1);
            border-radius: 14px;
            padding: 22px;
            box-shadow: 0 4px 16px rgba(0, 0, 0, 0.06), inset 0 1px 0 rgba(255, 255, 255, 0.8);
            margin-bottom: 16px;
            transition: all 0.35s cubic-bezier(0.4, 0, 0.2, 1);
        }

        .custom-card:hover {
            background: rgba(255, 255, 255, 0.95);
            border-color: rgba(201, 162, 39, 0.35);
            box-shadow: 0 8px 24px rgba(0, 0, 0, 0.08), 0 0 16px rgba(201, 162, 39, 0.06), inset 0 1px 0 rgba(255, 255, 255, 0.9);
            transform: translateY(-1px);
        }

        .custom-card h2 {
            margin-top: 0 !important;
            color: #1a3a52 !important;
        }

        .small-note {
            color: #6b7280 !important;
            font-size: 13px !important;
            line-height: 1.65 !important;
            font-weight: 400 !important;
        }

        /* --- Glassmorphism Section Cards --- */
        .section-card {
            background: rgba(255, 255, 255, 0.85);
            backdrop-filter: blur(16px) saturate(1.3);
            -webkit-backdrop-filter: blur(16px) saturate(1.3);
            border: 1px solid rgba(26, 38, 58, 0.1);
            border-radius: 14px;
            padding: 22px;
            box-shadow: 0 4px 16px rgba(0, 0, 0, 0.06), inset 0 1px 0 rgba(255, 255, 255, 0.8);
            margin-top: 20px;
            margin-bottom: 18px;
            transition: all 0.35s cubic-bezier(0.4, 0, 0.2, 1);
        }

        .section-card:hover {
            background: rgba(255, 255, 255, 0.95);
            border-color: rgba(201, 162, 39, 0.35);
            box-shadow: 0 8px 24px rgba(0, 0, 0, 0.08), 0 0 16px rgba(201, 162, 39, 0.06), inset 0 1px 0 rgba(255, 255, 255, 0.9);
            transform: translateY(-1px);
        }

        .section-card h2 {
            margin-top: 0 !important;
            color: #1a3a52 !important;
        }

        .criteria-list {
            color: #6b7280 !important;
            font-size: 13px !important;
            line-height: 1.7 !important;
            font-weight: 400 !important;
        }

        /* --- Glassmorphism File Uploader --- */
        [data-testid="stFileUploader"] {
            background: rgba(249, 250, 251, 0.85) !important;
            backdrop-filter: blur(12px) saturate(1.2) !important;
            -webkit-backdrop-filter: blur(12px) saturate(1.2) !important;
            border: 1px solid rgba(26, 38, 58, 0.1) !important;
            border-radius: 12px !important;
            padding: 14px !important;
            box-shadow: 0 2px 8px rgba(0, 0, 0, 0.04);
            transition: all 0.35s cubic-bezier(0.4, 0, 0.2, 1);
        }

        [data-testid="stFileUploader"]:hover {
            border-color: rgba(201, 162, 39, 0.35) !important;
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.06), 0 0 10px rgba(201, 162, 39, 0.04);
            background: rgba(255, 255, 255, 0.9) !important;
        }

        [data-testid="stFileUploaderDropzone"] {
            background: rgba(249, 250, 251, 0.6) !important;
            border: 2px dashed rgba(26, 58, 82, 0.25) !important;
            border-radius: 10px !important;
            color: #374151 !important;
            padding: 16px !important;
            transition: all 0.3s ease;
        }

        [data-testid="stFileUploaderDropzone"]:hover {
            border-color: #1a3a52 !important;
            background: rgba(243, 244, 246, 0.8) !important;
        }

        [data-testid="stFileUploaderDropzone"] * {
            color: #374151 !important;
            font-family: 'Montserrat', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
            font-size: 13px !important;
        }

        [data-testid="stFileUploaderDropzone"] svg {
            fill: #1a3a52 !important;
            color: #1a3a52 !important;
        }

        [data-testid="stFileUploaderDropzone"] button {
            background-color: #1a3a52 !important;
            color: #ffffff !important;
            border: none !important;
            border-radius: 8px !important;
            font-weight: 600 !important;
            font-size: 13px !important;
            padding: 0.5rem 1rem !important;
            box-shadow: 0 2px 8px rgba(26, 58, 82, 0.2);
            transition: all 0.2s ease;
        }

        [data-testid="stFileUploaderDropzone"] button * {
            color: #ffffff !important;
            font-size: 13px !important;
        }

        [data-testid="stFileUploaderDropzone"] button:hover {
            background-color: #12304a !important;
            color: #ffffff !important;
            box-shadow: 0 4px 12px rgba(26, 58, 82, 0.3);
        }

        /* --- Glassmorphism Buttons --- */
        .stButton > button,
        .stDownloadButton > button {
            background: rgba(26, 58, 82, 0.9) !important;
            backdrop-filter: blur(8px) !important;
            -webkit-backdrop-filter: blur(8px) !important;
            color: #ffffff !important;
            border: 1px solid rgba(201, 162, 39, 0.4) !important;
            border-radius: 9px !important;
            padding: 0.55rem 1.2rem !important;
            font-weight: 600 !important;
            font-size: 13px !important;
            font-family: 'Montserrat', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
            box-shadow: 0 2px 10px rgba(26, 58, 82, 0.15), inset 0 1px 0 rgba(255, 255, 255, 0.1);
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        }

        .stButton > button *,
        .stDownloadButton > button * {
            color: #ffffff !important;
            font-size: 13px !important;
        }

        .stButton > button:hover,
        .stDownloadButton > button:hover {
            background: rgba(26, 58, 82, 1) !important;
            border-color: rgba(201, 162, 39, 0.6) !important;
            color: #ffffff !important;
            transform: translateY(-1px);
            box-shadow: 0 4px 16px rgba(26, 58, 82, 0.25), 0 0 14px rgba(201, 162, 39, 0.1), inset 0 1px 0 rgba(255, 255, 255, 0.15);
        }

        /* --- Disabled buttons --- */
        .stButton > button:disabled,
        .stDownloadButton > button:disabled {
            background: rgba(229, 231, 235, 0.6) !important;
            backdrop-filter: blur(8px) !important;
            border: 1px solid rgba(209, 213, 219, 0.5) !important;
            color: rgba(107, 114, 128, 0.6) !important;
            box-shadow: 0 1px 4px rgba(0, 0, 0, 0.05) !important;
            cursor: not-allowed !important;
            opacity: 0.7 !important;
        }

        .stButton > button:disabled *,
        .stDownloadButton > button:disabled * {
            color: rgba(107, 114, 128, 0.6) !important;
        }

        [data-testid="stMarkdownContainer"] {
            color: #374151 !important;
            font-size: 14px !important;
        }

        .stCaption {
            font-size: 12px !important;
            color: #9ca3af !important;
        }

        /* --- Glassmorphism Input Fields --- */
        [data-testid="stTextInput"] input,
        [data-testid="stTextArea"] textarea {
            background: rgba(249, 250, 251, 0.9) !important;
            backdrop-filter: blur(10px) !important;
            -webkit-backdrop-filter: blur(10px) !important;
            color: #1e293b !important;
            border: 1px solid rgba(26, 38, 58, 0.15) !important;
            border-radius: 8px !important;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        }

        [data-testid="stTextInput"] input:hover,
        [data-testid="stTextArea"] textarea:hover {
            border-color: rgba(201, 162, 39, 0.4) !important;
            background: rgba(255, 255, 255, 0.95) !important;
        }

        [data-testid="stTextInput"] input:focus,
        [data-testid="stTextArea"] textarea:focus {
            border-color: #1a3a52 !important;
            box-shadow: 0 0 0 2px rgba(26, 58, 82, 0.1), 0 2px 8px rgba(0, 0, 0, 0.06) !important;
        }

        [data-testid="stSidebar"] {
            background: rgba(249, 250, 251, 0.9) !important;
            backdrop-filter: blur(20px) saturate(1.3) !important;
            -webkit-backdrop-filter: blur(20px) saturate(1.3) !important;
            border-right: 1px solid rgba(26, 38, 58, 0.08) !important;
        }

        /* ===== Override ALL Streamlit alert backgrounds ===== */
        [data-testid="stAlert"],
        [data-testid="stAlert-success"],
        [data-testid="stAlert-error"],
        [data-testid="stAlert-warning"],
        [data-testid="stAlert-info"],
        [data-testid="stStatusWidget"] {
            background: rgba(249, 250, 251, 0.9) !important;
            backdrop-filter: blur(14px) saturate(1.2) !important;
            -webkit-backdrop-filter: blur(14px) saturate(1.2) !important;
            border: 1px solid rgba(26, 38, 58, 0.1) !important;
            color: #374151 !important;
            box-shadow: 0 2px 10px rgba(0, 0, 0, 0.04) !important;
            transition: all 0.35s cubic-bezier(0.4, 0, 0.2, 1);
        }

        [data-testid="stAlert"]:hover,
        [data-testid="stAlert-success"]:hover,
        [data-testid="stAlert-error"]:hover,
        [data-testid="stAlert-warning"]:hover,
        [data-testid="stAlert-info"]:hover {
            background: rgba(255, 255, 255, 0.95) !important;
            border-color: rgba(201, 162, 39, 0.3) !important;
        }

        /* Force ALL alert icons to gold */
        [data-testid="stAlert"] svg,
        [data-testid="stAlert-success"] svg,
        [data-testid="stAlert-info"] svg,
        [data-testid="stAlert-warning"] svg,
        [data-testid="stAlert-error"] svg,
        [data-testid="stAlert"] [data-testid="stIcon"],
        [data-testid="stAlert-success"] [data-testid="stIcon"],
        [data-testid="stAlert-info"] [data-testid="stIcon"],
        [data-testid="stAlert-warning"] [data-testid="stIcon"],
        [data-testid="stAlert-error"] [data-testid="stIcon"] {
            fill: #c9a227 !important;
            color: #c9a227 !important;
        }

        [data-testid="stAlert-error"] svg,
        [data-testid="stAlert-error"] [data-testid="stIcon"] {
            fill: #dc2626 !important;
            color: #dc2626 !important;
        }

        [data-testid="stAlert-warning"] svg,
        [data-testid="stAlert-warning"] [data-testid="stIcon"] {
            fill: #ca8a04 !important;
            color: #ca8a04 !important;
        }

        /* Force all alert inner elements to use our text color */
        [data-testid="stAlert"] div,
        [data-testid="stAlert"] span,
        [data-testid="stAlert"] p,
        [data-testid="stAlert"] label {
            color: #374151 !important;
        }

        /* --- Custom glassmorphism alert divs (used via HTML) --- */
        .glass-alert {
            background: rgba(249, 250, 251, 0.9);
            backdrop-filter: blur(14px) saturate(1.2);
            -webkit-backdrop-filter: blur(14px) saturate(1.2);
            border-radius: 10px;
            padding: 12px 18px;
            margin-bottom: 12px;
            color: #374151 !important;
            font-family: 'Montserrat', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
            font-size: 13px;
            line-height: 1.6;
            box-shadow: 0 2px 10px rgba(0, 0, 0, 0.04);
            transition: all 0.35s cubic-bezier(0.4, 0, 0.2, 1);
            display: flex;
            align-items: center;
            gap: 8px;
        }

        .glass-alert-icon {
            font-size: 16px !important;
            font-weight: 700 !important;
            flex-shrink: 0;
            display: inline-flex;
            align-items: center;
            justify-content: center;
            width: 22px;
            height: 22px;
            border-radius: 50%;
        }

        .glass-alert:hover {
            background: rgba(255, 255, 255, 0.95);
        }

        .glass-alert-success {
            border: 1px solid rgba(201, 162, 39, 0.3);
        }
        .glass-alert-success .glass-alert-icon {
            color: #c9a227 !important;
            background: rgba(201, 162, 39, 0.1) !important;
        }

        .glass-alert-info {
            border: 1px solid rgba(26, 58, 82, 0.15);
        }
        .glass-alert-info .glass-alert-icon {
            color: #1a3a52 !important;
            background: rgba(26, 58, 82, 0.08) !important;
        }

        .glass-alert-warning {
            border: 1px solid rgba(202, 138, 4, 0.3);
        }
        .glass-alert-warning .glass-alert-icon {
            color: #ca8a04 !important;
            background: rgba(202, 138, 4, 0.1) !important;
        }

        .glass-alert-error {
            border: 1px solid rgba(220, 38, 38, 0.3);
        }
        .glass-alert-error .glass-alert-icon {
            color: #dc2626 !important;
            background: rgba(220, 38, 38, 0.1) !important;
        }

        /* --- Override Streamlit default blue for progress bars, etc. --- */
        .stProgress > div > div > div {
            background-color: #1a3a52 !important;
        }
        [data-testid="stTooltip"] {
            background: rgba(255, 255, 255, 0.95) !important;
            backdrop-filter: blur(10px) !important;
            border: 1px solid rgba(26, 38, 58, 0.1) !important;
        }

        /* --- Override ALL remaining Streamlit blue accents --- */
        [class*="stAlert"] [class*="icon"],
        [class*="stAlert"] svg,
        .stAlert svg {
            fill: #c9a227 !important;
            color: #c9a227 !important;
        }

        /* Force Streamlit spinner to navy */
        .stSpinner > div > svg {
            fill: #1a3a52 !important;
            color: #1a3a52 !important;
        }

        /* --- Scrollbar --- */
        ::-webkit-scrollbar {
            width: 8px;
        }
        ::-webkit-scrollbar-track {
            background: #f3f4f6;
        }
        ::-webkit-scrollbar-thumb {
            background: #d1d5db;
            border-radius: 4px;
            transition: background 0.2s ease;
        }
        ::-webkit-scrollbar-thumb:hover {
            background: #1a3a52;
        }

        /* --- Checkbox and radio --- */
        .stCheckbox label, .stRadio label {
            color: #374151 !important;
        }

        /* --- Glassmorphism Selectbox --- */
        [data-baseweb="select"] > div {
            background: rgba(249, 250, 251, 0.9) !important;
            backdrop-filter: blur(10px) !important;
            -webkit-backdrop-filter: blur(10px) !important;
            border-color: rgba(26, 38, 58, 0.15) !important;
            color: #1e293b !important;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        }

        [data-basewase="select"] > div:hover {
            border-color: rgba(201, 162, 39, 0.4) !important;
            background: rgba(255, 255, 255, 0.95) !important;
        }

        /* --- Tabs (navy accent, gold highlight) --- */
        [data-testid="stTabs"] button {
            color: #9ca3af !important;
            transition: all 0.3s ease;
        }
        [data-testid="stTabs"] button:hover {
            color: #c9a227 !important;
        }
        [data-testid="stTabs"] button[aria-selected="true"] {
            color: #1a3a52 !important;
            border-bottom-color: #1a3a52 !important;
        }

        /* --- Dividers --- */
        hr {
            border-color: rgba(26, 38, 58, 0.1);
            transition: border-color 0.3s ease;
        }
        hr:hover {
            border-color: rgba(201, 162, 39, 0.3);
        }

        /* --- Spinner --- */
        .stSpinner > div {
            border-color: rgba(26, 58, 82, 0.15) !important;
            border-top-color: #1a3a52 !important;
        }

        /* --- Override all Streamlit blue accents to navy/gold --- */
        [data-testid="stCaption"], .stCaption {
            color: #9ca3af !important;
        }

        /* --- Glassmorphism block container --- */
        .block-container {
            padding-top: 3rem;
            padding-bottom: 3rem;
            max-width: 1200px;
        }
        </style>
        """,
        unsafe_allow_html=True
    )


def init_session_state():
    defaults = {
        "written_report_content": None,
        "written_pdf_bytes": None,
        "written_pdf_file_name": None,
        "presentation_report_content": None,
        "presentation_pdf_bytes": None,
        "presentation_pdf_file_name": None,
        "combined_pdf_bytes": None,
        "combined_pdf_file_name": None,
        "mentor_pack_content": None,
        "mentor_pack_pdf_bytes": None,
        "mentor_pack_pdf_file_name": None,
        "mentor_name": None,
        "video_transcript": None,
        "video_summary": None,
    }

    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def save_uploaded_file(uploaded_file, folder: Path) -> str:
    folder.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_file_name = uploaded_file.name.replace(" ", "_")
    file_path = folder / f"{timestamp}_{safe_file_name}"

    with open(file_path, "wb") as file:
        file.write(uploaded_file.getbuffer())

    return str(file_path)


def read_markdown_file(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as file:
        return file.read()


def read_text_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
        return file.read()


def read_text_from_pdf(file_path: str) -> str:
    text = ""
    reader = PdfReader(file_path)

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"

    return text


def read_text_from_docx(file_path: str) -> str:
    document = Document(file_path)
    text_parts = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())

    return "\n".join(text_parts)


def read_uploaded_document_text(file_path: str) -> str:
    extension = Path(file_path).suffix.lower()

    if extension == ".txt":
        return read_text_from_txt(file_path)

    if extension == ".pdf":
        return read_text_from_pdf(file_path)

    if extension == ".docx":
        return read_text_from_docx(file_path)

    raise ValueError(f"Unsupported file type: {extension}")


def clean_detected_name(name: str) -> Optional[str]:
    if not name:
        return None

    name = name.strip()
    name = re.sub(r"\s+", " ", name)

    unwanted_words = [
        "student name",
        "learner name",
        "assessor name",
        "date issued",
        "completion date",
        "submitted on",
        "project title",
        "module name",
        "qualification name",
        "product name",
        "student signature",
        "learner declaration"
    ]

    lower_name = name.lower()

    for word in unwanted_words:
        if word in lower_name:
            return None

    if len(name) < 3 or len(name) > 80:
        return None

    if not re.search(r"[A-Za-z]", name):
        return None

    return name


def extract_student_name_from_docx_tables(file_path: str) -> Optional[str]:
    document = Document(file_path)

    for table in document.tables:
        rows = table.rows

        for row_index, row in enumerate(rows):
            cells = row.cells

            for cell_index, cell in enumerate(cells):
                cell_text = cell.text.strip().lower()

                if "student name" in cell_text or "learner name" in cell_text:
                    if row_index + 1 < len(rows):
                        possible_name = rows[row_index + 1].cells[cell_index].text.strip()
                        cleaned_name = clean_detected_name(possible_name)

                        if cleaned_name:
                            return cleaned_name

                    if cell_index + 1 < len(cells):
                        possible_name = cells[cell_index + 1].text.strip()
                        cleaned_name = clean_detected_name(possible_name)

                        if cleaned_name:
                            return cleaned_name

    return None


def extract_student_name_from_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    full_text = "\n".join(lines)

    patterns = [
        r"student\s*name\s*[:\-]\s*([A-Za-z][A-Za-z\s.'-]{2,80})",
        r"learner\s*name\s*[:\-]\s*([A-Za-z][A-Za-z\s.'-]{2,80})",
        r"name\s*of\s*student\s*[:\-]\s*([A-Za-z][A-Za-z\s.'-]{2,80})"
    ]

    for pattern in patterns:
        match = re.search(pattern, full_text, re.IGNORECASE)
        if match:
            cleaned_name = clean_detected_name(match.group(1))
            if cleaned_name:
                return cleaned_name

    for index, line in enumerate(lines):
        lower_line = line.lower()

        if "student name" in lower_line or "learner name" in lower_line:
            next_lines = lines[index + 1:index + 8]

            for possible_name in next_lines:
                cleaned_name = clean_detected_name(possible_name)

                if cleaned_name:
                    if not cleaned_name.lower().startswith(("mr.", "mrs.", "ms.", "dr.")):
                        return cleaned_name

    return "Unknown"



def extract_mentor_name_from_docx_tables(file_path: str) -> Optional[str]:
    """
    Detects the mentor / assessor name from DOCX tables.
    This supports common front-page layouts where the table has:
    Student name | Assessor name
    Learner name | Mentor name
    and the actual names are in the next row.
    """
    document = Document(file_path)

    assessor_labels = [
        "assessor name",
        "mentor name",
        "teacher name",
        "trainer name",
        "lecturer name",
        "tutor name"
    ]

    for table in document.tables:
        rows = table.rows

        for row_index, row in enumerate(rows):
            cells = row.cells

            for cell_index, cell in enumerate(cells):
                cell_text = cell.text.strip().lower()

                if any(label in cell_text for label in assessor_labels):
                    if row_index + 1 < len(rows):
                        possible_name = rows[row_index + 1].cells[cell_index].text.strip()
                        cleaned_name = clean_detected_name(possible_name)

                        if cleaned_name:
                            return cleaned_name

                    if cell_index + 1 < len(cells):
                        possible_name = cells[cell_index + 1].text.strip()
                        cleaned_name = clean_detected_name(possible_name)

                        if cleaned_name:
                            return cleaned_name

    return None


def extract_mentor_name_from_text(text: str) -> str:
    """
    Detects mentor / assessor name from plain text extracted from PDF, DOCX, or TXT.
    Table extraction is more accurate for DOCX, but this helps for PDFs.
    """
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    full_text = "\n".join(lines)

    patterns = [
        r"assessor\s*name\s*[:\-]\s*([A-Za-z][A-Za-z\s.'-]{2,80})",
        r"mentor\s*name\s*[:\-]\s*([A-Za-z][A-Za-z\s.'-]{2,80})",
        r"teacher\s*name\s*[:\-]\s*([A-Za-z][A-Za-z\s.'-]{2,80})",
        r"trainer\s*name\s*[:\-]\s*([A-Za-z][A-Za-z\s.'-]{2,80})",
        r"lecturer\s*name\s*[:\-]\s*([A-Za-z][A-Za-z\s.'-]{2,80})",
        r"tutor\s*name\s*[:\-]\s*([A-Za-z][A-Za-z\s.'-]{2,80})"
    ]

    for pattern in patterns:
        match = re.search(pattern, full_text, re.IGNORECASE)
        if match:
            cleaned_name = clean_detected_name(match.group(1))
            if cleaned_name:
                return cleaned_name

    assessor_labels = [
        "assessor name",
        "mentor name",
        "teacher name",
        "trainer name",
        "lecturer name",
        "tutor name"
    ]

    for index, line in enumerate(lines):
        lower_line = line.lower()

        if any(label in lower_line for label in assessor_labels):
            next_lines = lines[index + 1:index + 8]

            for possible_name in next_lines:
                parts = re.split(r"\s{2,}|\t+", possible_name.strip())
                if len(parts) >= 2:
                    possible_name = parts[-1].strip()

                cleaned_name = clean_detected_name(possible_name)

                if cleaned_name:
                    return cleaned_name

    return "Unknown"


def detect_mentor_name(file_path: str) -> str:
    """
    Detects the assessor / mentor name from the uploaded learner submission.
    In your report template, the Assessor name is treated as the mentor name.
    """
    file_path_obj = Path(file_path)
    extension = file_path_obj.suffix.lower()

    try:
        if extension == ".docx":
            table_name = extract_mentor_name_from_docx_tables(file_path)

            if table_name:
                return table_name

            text = read_text_from_docx(file_path)
            return extract_mentor_name_from_text(text)

        if extension == ".pdf":
            text = read_text_from_pdf(file_path)
            return extract_mentor_name_from_text(text)

        if extension == ".txt":
            text = read_text_from_txt(file_path)
            return extract_mentor_name_from_text(text)

        return "Unknown"

    except Exception:
        return "Unknown"


def detect_student_name(file_path: str) -> str:
    file_path_obj = Path(file_path)
    extension = file_path_obj.suffix.lower()

    try:
        if extension == ".docx":
            table_name = extract_student_name_from_docx_tables(file_path)

            if table_name:
                return table_name

            text = read_text_from_docx(file_path)
            return extract_student_name_from_text(text)

        if extension == ".pdf":
            text = read_text_from_pdf(file_path)
            return extract_student_name_from_text(text)

        if extension == ".txt":
            text = read_text_from_txt(file_path)
            return extract_student_name_from_text(text)

        return "Unknown"

    except Exception:
        return "Unknown"


def format_pdf_text(text: str) -> str:
    text = escape(text)

    text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"`(.*?)`", r"<font name='Courier'>\1</font>", text)

    return text


def convert_markdown_to_pdf(markdown_text: str, pdf_path: Path) -> str:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        rightMargin=45,
        leftMargin=45,
        topMargin=50,
        bottomMargin=50
    )

    styles = getSampleStyleSheet()

    # Professional, compact title style
    styles.add(
        ParagraphStyle(
            name="ProfTitle",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            spaceAfter=10,
            spaceBefore=6,
            textColor="#1a3a52",
            alignment=TA_LEFT,
        )
    )

    # Section headings (## Heading)
    styles.add(
        ParagraphStyle(
            name="ProfH2",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=16,
            spaceAfter=6,
            spaceBefore=14,
            textColor="#12304a",
            alignment=TA_LEFT,
        )
    )

    # Sub headings (### Heading)
    styles.add(
        ParagraphStyle(
            name="ProfH3",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=14,
            spaceAfter=4,
            spaceBefore=10,
            textColor="#1e293b",
            alignment=TA_LEFT,
        )
    )

    # Normal body text
    styles.add(
        ParagraphStyle(
            name="ProfNormal",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=13,
            spaceAfter=4,
            alignment=TA_LEFT,
            textColor="#1e293b",
        )
    )

    # Bullet / numbered list text
    styles.add(
        ParagraphStyle(
            name="ProfBullet",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=13,
            leftIndent=18,
            firstLineIndent=-10,
            spaceAfter=3,
            alignment=TA_LEFT,
            textColor="#1e293b",
        )
    )

    story = []

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()

        if not line:
            story.append(Spacer(1, 6))
            continue

        # Main title (# Heading)
        if line.startswith("# "):
            story.append(
                Paragraph(format_pdf_text(line[2:]), styles["ProfTitle"])
            )

        # Section headings (## Heading)
        elif line.startswith("## "):
            story.append(
                Paragraph(format_pdf_text(line[3:]), styles["ProfH2"])
            )

        # Sub headings (### Heading)
        elif line.startswith("### "):
            story.append(
                Paragraph(format_pdf_text(line[4:]), styles["ProfH3"])
            )

        # Presentation criteria headings
        elif re.match(
            r"^\d+\.\s+(Use of Visuals|Clarity\s*&\s*Structure|Key Points Coverage|Understanding|Language\s*&\s*Communication|Time Management|Original Thinking|Q&A Handling)",
            line,
            re.IGNORECASE
        ):
            story.append(
                Paragraph(format_pdf_text(line), styles["ProfH3"])
            )

        # Bullet points
        elif line.startswith("- ") or line.startswith("* "):
            story.append(
                Paragraph("• " + format_pdf_text(line[2:]), styles["ProfBullet"])
            )

        # Normal numbered points
        elif re.match(r"^\d+\.\s+", line):
            story.append(
                Paragraph(format_pdf_text(line), styles["ProfBullet"])
            )

        # Normal paragraph text
        else:
            story.append(
                Paragraph(format_pdf_text(line), styles["ProfNormal"])
            )

    doc.build(story)

    return str(pdf_path)


def replace_unknown_learner_name(report_content: str, learner_name: str) -> str:
    if learner_name == "Unknown":
        return report_content

    replacements = {
        "Learner Name: Unknown": f"Learner Name: {learner_name}",
        "Learner name: Unknown": f"Learner name: {learner_name}",
        "Student Name: Unknown": f"Student Name: {learner_name}",
        "Student name: Unknown": f"Student name: {learner_name}",
        "**Learner Name:** Unknown": f"**Learner Name:** {learner_name}",
        "**Student Name:** Unknown": f"**Student Name:** {learner_name}",
    }

    for old_text, new_text in replacements.items():
        report_content = report_content.replace(old_text, new_text)

    return report_content


def save_markdown_report(content: str, output_file_name: str) -> str:
    REPORT_FOLDER.mkdir(parents=True, exist_ok=True)
    output_path = REPORT_FOLDER / output_file_name

    with open(output_path, "w", encoding="utf-8") as file:
        file.write(content)

    return str(output_path)



def create_combined_feedback_pdf(written_feedback: str, presentation_feedback: str):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    combined_markdown = f"""
# Combined Feedback Report

## Written Report Feedback

{written_feedback}

---

## Presentation Feedback

{presentation_feedback}
"""

    combined_pdf_file_name = f"combined_feedback_report_and_presentation_{timestamp}.pdf"
    combined_pdf_path = REPORT_FOLDER / combined_pdf_file_name

    convert_markdown_to_pdf(
        markdown_text=combined_markdown,
        pdf_path=combined_pdf_path
    )

    with open(combined_pdf_path, "rb") as pdf_file:
        combined_pdf_bytes = pdf_file.read()

    return combined_pdf_bytes, combined_pdf_file_name



def create_mentor_pack_pdf(mentor_pack_content: str):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    mentor_pack_pdf_file_name = f"mentor_feedback_pack_{timestamp}.pdf"
    mentor_pack_pdf_path = REPORT_FOLDER / mentor_pack_pdf_file_name

    convert_markdown_to_pdf(
        markdown_text=mentor_pack_content,
        pdf_path=mentor_pack_pdf_path
    )

    with open(mentor_pack_pdf_path, "rb") as pdf_file:
        mentor_pack_pdf_bytes = pdf_file.read()

    return mentor_pack_pdf_bytes, mentor_pack_pdf_file_name

def main():
    import base64

    LOGO_PATH = Path(__file__).parent / "assets" / "educlaas_logo.png"
    logo_exists = LOGO_PATH.exists()

    st.set_page_config(
        page_title="Learner Submission Feedback Agent",
        page_icon=str(LOGO_PATH) if logo_exists else "📝",
        layout="wide"
    )

    apply_custom_css()
    init_session_state()

    if logo_exists:
        logo_b64 = base64.b64encode(LOGO_PATH.read_bytes()).decode()
        st.markdown(
            f'<div style="display:flex;align-items:center;gap:12px;margin-bottom:0.6rem;padding-right:30px;padding-top:30px;padding-bottom:30px;">'
            f'<img src="data:image/png;base64,{logo_b64}" style="height:46px;width:auto;border-radius:8px;object-fit:contain;" />'
            f'<h1 style="margin:0;padding-left:20px;color:#1a3a52;font-family:\'Montserrat\',sans-serif;font-weight:600;font-size:1.6rem;letter-spacing:-0.3px;">'
            f'Learner Submission Feedback Agent'
            f'</h1>'
            f'</div>',
            unsafe_allow_html=True
        )
    else:
        st.title("📝 Learner Submission Feedback Agent")

    # ============================================================
    # SECTION 1: Project Report Feedback
    # ============================================================
    st.markdown(
        '<div class="custom-card">'
        '<h2>Project Report Feedback</h2>'
        '<p class="small-note">'
        'Upload the project brief and learner report to generate a clear, '
        'professional feedback. The system can detect the learner name '
        'and create a downloadable PDF report.'
        '</p>'
        '</div>',
        unsafe_allow_html=True
    )

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Upload project brief")
        project_brief_file = st.file_uploader(
            "Upload project brief",
            type=["pdf", "docx", "txt"],
            key="brief_upload"
        )

    with col2:
        st.subheader("Upload learner submission")
        submission_file = st.file_uploader(
            "Upload learner submission",
            type=["pdf", "docx", "txt"],
            key="submission_upload"
        )


    if st.button("Generate Feedback Report", key="btn_generate_report"):
        if run_feedback_pipeline is None:
            st.markdown(
                '<div class="glass-alert glass-alert-error"><span class="glass-alert-icon">✗</span> Feedback pipeline backend is missing. Please check that <code>backend/pipeline.py</code> exists.</div>',
                unsafe_allow_html=True
            )
            return

        if project_brief_file is None:
            st.markdown('<div class="glass-alert glass-alert-error"><span class="glass-alert-icon">✗</span> Please upload the project brief first.</div>', unsafe_allow_html=True)
            return

        if submission_file is None:
            st.markdown('<div class="glass-alert glass-alert-error"><span class="glass-alert-icon">✗</span> Please upload the learner submission first.</div>', unsafe_allow_html=True)
            return

        with st.spinner("Reading files and generating written feedback..."):
            try:
                brief_path = save_uploaded_file(project_brief_file, BRIEF_FOLDER)
                submission_path = save_uploaded_file(submission_file, SUBMISSION_FOLDER)

                learner_name = detect_student_name(submission_path)
                mentor_name = detect_mentor_name(submission_path)

                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                markdown_file_name = f"feedback_report_{timestamp}.md"

                report_path = run_feedback_pipeline(
                    brief_file_path=brief_path,
                    submission_file_path=submission_path,
                    output_folder=str(REPORT_FOLDER),
                    output_file_name=markdown_file_name
                )

                report_content = read_markdown_file(report_path)

                report_content = replace_unknown_learner_name(
                    report_content=report_content,
                    learner_name=learner_name
                )

                with open(report_path, "w", encoding="utf-8") as file:
                    file.write(report_content)

                pdf_file_name = f"feedback_report_{timestamp}.pdf"
                pdf_path = REPORT_FOLDER / pdf_file_name

                convert_markdown_to_pdf(
                    markdown_text=report_content,
                    pdf_path=pdf_path
                )

                with open(pdf_path, "rb") as pdf_file:
                    pdf_bytes = pdf_file.read()

                st.session_state.written_report_content = report_content
                st.session_state.written_pdf_bytes = pdf_bytes
                st.session_state.written_pdf_file_name = pdf_file_name
                st.session_state.mentor_name = mentor_name

                st.markdown('<div class="glass-alert glass-alert-success"><span class="glass-alert-icon">✓</span> Written feedback report generated successfully!</div>', unsafe_allow_html=True)

                if learner_name != "Unknown":
                    st.markdown(f'<div class="glass-alert glass-alert-info"><span class="glass-alert-icon">ℹ</span> Learner name detected: {learner_name}</div>', unsafe_allow_html=True)
                else:
                    st.markdown('<div class="glass-alert glass-alert-warning"><span class="glass-alert-icon">⚠</span> Learner name could not be detected automatically.</div>', unsafe_allow_html=True)

                if mentor_name != "Unknown":
                    st.markdown(f'<div class="glass-alert glass-alert-info"><span class="glass-alert-icon">ℹ</span> Mentor / assessor name detected: {mentor_name}</div>', unsafe_allow_html=True)
                else:
                    st.markdown('<div class="glass-alert glass-alert-warning"><span class="glass-alert-icon">⚠</span> Mentor / assessor name could not be detected automatically.</div>', unsafe_allow_html=True)

            except Exception as error:
                st.markdown('<div class="glass-alert glass-alert-error"><span class="glass-alert-icon">✗</span> Something went wrong while generating written feedback.</div>', unsafe_allow_html=True)
                st.exception(error)

    if st.session_state.written_report_content:
        st.subheader("Generated Written Feedback Report")
        st.markdown(st.session_state.written_report_content)

        st.download_button(
            label="Download Written Feedback Report as PDF",
            data=st.session_state.written_pdf_bytes,
            file_name=st.session_state.written_pdf_file_name,
            mime="application/pdf"
        )

        st.divider()

    st.divider()    

    # ============================================================
    # SECTION 2: Presentation Feedback
    # ============================================================
    st.markdown(
        '<div class="section-card">'
        '<h2>Presentation Feedback</h2>'
        '<p class="small-note">'
        'Upload the Teams presentation transcript to generate feedback '
        'only for the learner presentation. This will create a presentation '
        'feedback PDF.'
        '</p>'
        '<p class="criteria-list">'
        'Assessment criteria: Use of Visuals, Clarity &amp; Structure, Key Points Coverage, '
        'Understanding, Language &amp; Communication, Time Management, Original Thinking, '
        'and Q&amp;A Handling.'
        '</p>'
        '</div>',
        unsafe_allow_html=True
    )

    st.subheader("Upload presentation transcript")

    presentation_transcript_file = st.file_uploader(
        "Upload presentation transcript",
        type=["pdf", "docx", "txt"],
        key="presentation_upload"
    )

    if st.button("Generate Presentation Feedback", key="btn_generate_presentation"):
        if evaluate_presentation is None:
            st.markdown(
                '<div class="glass-alert glass-alert-error"><span class="glass-alert-icon">✗</span> Presentation evaluator backend is missing. Please check that <code>backend/presentation_evaluator.py</code> exists.</div>',
                unsafe_allow_html=True
            )
            return

        if presentation_transcript_file is None:
            st.markdown('<div class="glass-alert glass-alert-error"><span class="glass-alert-icon">✗</span> Please upload the presentation transcript first.</div>', unsafe_allow_html=True)
            return

        with st.spinner("Reading transcript and generating presentation feedback..."):
            try:
                presentation_path = save_uploaded_file(
                    presentation_transcript_file,
                    PRESENTATION_FOLDER
                )

                transcript_text = read_uploaded_document_text(presentation_path)

                presentation_feedback = evaluate_presentation(transcript_text)

                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                presentation_markdown_file_name = f"presentation_feedback_{timestamp}.md"
                presentation_pdf_file_name = f"presentation_feedback_{timestamp}.pdf"

                presentation_report_path = save_markdown_report(
                    content=presentation_feedback,
                    output_file_name=presentation_markdown_file_name
                )

                presentation_pdf_path = REPORT_FOLDER / presentation_pdf_file_name

                convert_markdown_to_pdf(
                    markdown_text=presentation_feedback,
                    pdf_path=presentation_pdf_path
                )

                with open(presentation_pdf_path, "rb") as pdf_file:
                    presentation_pdf_bytes = pdf_file.read()

                st.session_state.presentation_report_content = presentation_feedback
                st.session_state.presentation_pdf_bytes = presentation_pdf_bytes
                st.session_state.presentation_pdf_file_name = presentation_pdf_file_name

                st.markdown('<div class="glass-alert glass-alert-success"><span class="glass-alert-icon">✓</span> Presentation feedback generated successfully!</div>', unsafe_allow_html=True)

            except Exception as error:
                st.markdown('<div class="glass-alert glass-alert-error"><span class="glass-alert-icon">✗</span> Something went wrong while generating presentation feedback.</div>', unsafe_allow_html=True)
                st.exception(error)

    if st.session_state.presentation_report_content:
        st.subheader("Generated Presentation Feedback")
        st.markdown(st.session_state.presentation_report_content)

        st.download_button(
            label="Download Presentation Feedback as PDF",
            data=st.session_state.presentation_pdf_bytes,
            file_name=st.session_state.presentation_pdf_file_name,
            mime="application/pdf"
        )



    st.divider()

    # ============================================================
    # SECTION 3: Combined Feedback PDF
    # ============================================================
    st.subheader("Combined Feedback PDF")

    written_ready = st.session_state.get("written_report_content") is not None
    presentation_ready = st.session_state.get("presentation_report_content") is not None

    st.caption(
        f"Written feedback detected: {'Yes' if written_ready else 'No'} | "
        f"Presentation feedback detected: {'Yes' if presentation_ready else 'No'}"
    )

    if written_ready and presentation_ready:
        combined_pdf_bytes, combined_pdf_file_name = create_combined_feedback_pdf(
            written_feedback=st.session_state.written_report_content,
            presentation_feedback=st.session_state.presentation_report_content
        )

        st.session_state.combined_pdf_bytes = combined_pdf_bytes
        st.session_state.combined_pdf_file_name = combined_pdf_file_name

        st.download_button(
            label="Download Combined Feedback for Report and Presentation as PDF",
            data=st.session_state.combined_pdf_bytes,
            file_name=st.session_state.combined_pdf_file_name,
            mime="application/pdf"
        )

    else:
        st.download_button(
            label="Download Combined Feedback for Report and Presentation as PDF",
            data=b"",
            file_name="combined_feedback_report_and_presentation.pdf",
            mime="application/pdf",
            disabled=True
        )

        st.markdown(
            '<div class="glass-alert glass-alert-info"><span class="glass-alert-icon">ℹ</span> Please generate both the written feedback report and the presentation feedback first. After both are generated, this button will become active.</div>',
            unsafe_allow_html=True
        )



    st.divider()

    # ============================================================
    # SECTION 4: Full Mentor Feedback Pack
    # ============================================================
    st.markdown(
        '<div class="section-card">'
        '<h2>Full Mentor Feedback Pack</h2>'
        '<p class="small-note">'
        'This creates a mentor-only pack using the feedback generated for the report and '
        'the presentation. It includes a mentor summary, missing '
        'requirements checklist, action plan, suggested marking table, '
        'learner-friendly message, and follow-up questions.'
        '</p>'
        '</div>',
        unsafe_allow_html=True
    )

    mentor_tools_ready = generate_mentor_feedback_pack is not None

    st.caption(
        f"Written feedback detected: {'Yes' if written_ready else 'No'} | "
        f"Presentation feedback detected: {'Yes' if presentation_ready else 'No'} | "
        f"Mentor evaluator detected: {'Yes' if mentor_tools_ready else 'No'} | "
        f"Mentor name: {st.session_state.get('mentor_name') or 'Not detected yet'}"
    )

    if not mentor_tools_ready:
        st.markdown(
            '<div class="glass-alert glass-alert-error"><span class="glass-alert-icon">✗</span> Mentor evaluator backend is missing. Please add <code>backend/mentor_evaluator.py</code> to your project.</div>',
            unsafe_allow_html=True
        )

    if written_ready and presentation_ready and mentor_tools_ready:
        if st.button("Generate Full Mentor Feedback Pack", key="btn_generate_mentor"):
            with st.spinner("Generating mentor summary, checklist, action plan and marking table..."):
                try:
                    mentor_pack_content = generate_mentor_feedback_pack(
                        written_feedback=st.session_state.written_report_content,
                        presentation_feedback=st.session_state.presentation_report_content,
                        mentor_name=st.session_state.get("mentor_name") or "Unknown"
                    )

                    mentor_pack_pdf_bytes, mentor_pack_pdf_file_name = create_mentor_pack_pdf(
                        mentor_pack_content=mentor_pack_content
                    )

                    st.session_state.mentor_pack_content = mentor_pack_content
                    st.session_state.mentor_pack_pdf_bytes = mentor_pack_pdf_bytes
                    st.session_state.mentor_pack_pdf_file_name = mentor_pack_pdf_file_name

                    st.markdown('<div class="glass-alert glass-alert-success"><span class="glass-alert-icon">✓</span> Full mentor feedback pack generated successfully!</div>', unsafe_allow_html=True)

                except Exception as error:
                    st.markdown('<div class="glass-alert glass-alert-error"><span class="glass-alert-icon">✗</span> Something went wrong while generating the mentor feedback pack.</div>', unsafe_allow_html=True)
                    st.exception(error)
    else:
        st.button(
            "Generate Full Mentor Feedback Pack",
            disabled=True,
            key="btn_generate_mentor_disabled"
        )

        st.markdown(
            '<div class="glass-alert glass-alert-info"><span class="glass-alert-icon">ℹ</span> Generate both the written feedback report and presentation feedback first. Then this button will become active.</div>',
            unsafe_allow_html=True
        )

    if st.session_state.mentor_pack_content:
        st.subheader("Generated Full Mentor Feedback Pack")
        st.markdown(st.session_state.mentor_pack_content)

        st.download_button(
            label="Download Full Mentor Feedback Pack as PDF",
            data=st.session_state.mentor_pack_pdf_bytes,
            file_name=st.session_state.mentor_pack_pdf_file_name,
            mime="application/pdf"
        )

    # ============================================================
    # SECTION 5: Video Transcription & Summarization
    # ============================================================
    st.markdown(
        '<div class="section-card">'
        '<h2>Video Transcription &amp; Summarization</h2>'
        '<p class="small-note">'
        'Upload a video file to extract the audio, transcribe the speech, '
        'and generate a detailed summary of the content.'
        '</p>'
        '</div>',
        unsafe_allow_html=True
    )

    video_backend_ready = (
        extract_audio is not None
        and transcribe_audio is not None
        and summarize_text is not None
        and chunked_summarize is not None
    )

    if not video_backend_ready:
        st.markdown(
            '<div class="glass-alert glass-alert-error"><span class="glass-alert-icon">✗</span> Video transcription aborted.Instance failed: 4rsrmRan out of memory (used over 512MB) while running your code.',
            unsafe_allow_html=True
        )

    uploaded_file = st.file_uploader(
        "Upload a video file",
        type=["mp4", "avi", "mkv", "mov"],
        key="video_upload"
    )

    if uploaded_file is not None:

        video_path = "uploaded_video.mp4"
        audio_path = "temp_audio.wav"
        transcript_path = "transcript.txt"
        summary_path = "summary.txt"

        with open(video_path, "wb") as f:
            f.write(uploaded_file.read())

        st.markdown('<div class="glass-alert glass-alert-success"><span class="glass-alert-icon">✓</span> Video uploaded successfully.</div>', unsafe_allow_html=True)

        # =========================
        # TRANSCRIBE BUTTON
        # =========================
        if st.button("Transcribe", key="btn_transcribe"):
            if not video_backend_ready:
                st.markdown(
                    '<div class="glass-alert glass-alert-error"><span class="glass-alert-icon">✗</span> Video transcription backend is not available.</div>',
                    unsafe_allow_html=True
                )
            else:
                extract_message = st.empty()
                extract_message.write("Extracting audio...")
                extract_audio(video_path, audio_path)
                extract_message.empty()

                transcribe_message = st.empty()
                transcribe_message.write("Transcribing audio...")
                transcript = transcribe_audio(
                    audio_path,
                    model_size="base"
                )
                transcribe_message.empty()

                with open(transcript_path, "w", encoding="utf-8") as f:
                    f.write(transcript)

                st.session_state.video_transcript = transcript

                st.subheader("Transcript")
                st.write(transcript)

                st.markdown('<div class="glass-alert glass-alert-success"><span class="glass-alert-icon">✓</span> Transcript saved as transcript.txt</div>', unsafe_allow_html=True)

        # =========================
        # SUMMARIZE BUTTON
        # =========================
        if st.button("Summarize", key="btn_summarize"):
            if not video_backend_ready:
                st.markdown(
                    '<div class="glass-alert glass-alert-error"><span class="glass-alert-icon">✗</span> Video summarization backend is not available.</div>',
                    unsafe_allow_html=True
                )
            elif not os.path.exists(transcript_path) and st.session_state.video_transcript is None:
                st.markdown('<div class="glass-alert glass-alert-warning"><span class="glass-alert-icon">⚠</span> Please click Transcribe first.</div>', unsafe_allow_html=True)
            else:
                if os.path.exists(transcript_path):
                    with open(transcript_path, "r", encoding="utf-8") as f:
                        transcript = f.read()
                else:
                    transcript = st.session_state.video_transcript

                summary_message = st.empty()
                summary_message.write("Generating detailed summary...")

                summary = chunked_summarize(
                    text=transcript,
                    summarize_func=lambda txt: summarize_text(
                        txt,
                        max_length=350,
                        min_length=120
                    ),
                    max_chunk_size=2000
                )

                summary_message.empty()

                with open(summary_path, "w", encoding="utf-8") as f:
                    f.write(summary)

                st.session_state.video_summary = summary

                st.subheader("Detailed Summary")
                st.write(summary)

                st.markdown('<div class="glass-alert glass-alert-success"><span class="glass-alert-icon">✓</span> Summary saved as summary.txt</div>', unsafe_allow_html=True)

    else:
        st.markdown('<div class="glass-alert glass-alert-info"><span class="glass-alert-icon">ℹ</span> Please upload a video file first.</div>', unsafe_allow_html=True)

if __name__ == "__main__":
    main()
