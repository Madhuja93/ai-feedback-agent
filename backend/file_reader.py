from pathlib import Path
from docx import Document
from pypdf import PdfReader


def read_txt_file(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
        return file.read()


def read_pdf_file(file_path: str) -> str:
    text = ""

    reader = PdfReader(file_path)

    for page in reader.pages:
        page_text = page.extract_text()

        if page_text:
            text += page_text + "\n"

    return text


def read_docx_file(file_path: str) -> str:
    document = Document(file_path)
    text_parts = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())

    return "\n".join(text_parts)


def read_file_text(file_path: str) -> str:
    extension = Path(file_path).suffix.lower()

    if extension == ".txt":
        return read_txt_file(file_path)

    if extension == ".pdf":
        return read_pdf_file(file_path)

    if extension == ".docx":
        return read_docx_file(file_path)

    raise ValueError(f"Unsupported file type: {extension}")